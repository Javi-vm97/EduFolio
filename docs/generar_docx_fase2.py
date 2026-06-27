# -*- coding: utf-8 -*-
"""
EduFolio - Generador de la documentacion de la Fase 2 (.docx)

Documento acumulativo: incluye los fundamentos del proyecto y el desarrollo
de la Segunda Fase (gestion completa de contenidos y manejo de archivos).

Formato (lineamientos institucionales):
  - Fuente Arial: 12 texto, 13 subtitulos, 14 titulos, 10 rotulos de tablas
  - Margenes: superior/inferior 2.5 cm, izquierdo 3.5 cm, derecho 2.5 cm
  - Interlineado sencillo, espacio entre parrafos, justificado, sangria 1 cm
  - Paginacion: romano minuscula en preliminares, arabigo en el cuerpo
Citas y referencias en estilo APA (7.a edicion).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NEGRO = RGBColor(0, 0, 0)


def _arial(run, size, bold=False, italic=False):
    run.font.name = 'Arial'; run.font.size = Pt(size)
    run.font.bold = bold; run.font.italic = italic; run.font.color.rgb = NEGRO
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rf.set(qn(a), 'Arial')


def page_field(p):
    r = p.add_run(); _arial(r, 12)
    for k, t in (('begin', None), (None, 'PAGE'), ('end', None)):
        if k:
            f = OxmlElement('w:fldChar'); f.set(qn('w:fldCharType'), k); r._r.append(f)
        else:
            it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = t; r._r.append(it)


def numbering(section, fmt, start=None):
    sp = section._sectPr; pg = sp.find(qn('w:pgNumType'))
    if pg is None:
        pg = OxmlElement('w:pgNumType'); sp.append(pg)
    pg.set(qn('w:fmt'), fmt)
    if start is not None:
        pg.set(qn('w:start'), str(start))


def footer_num(section):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    page_field(p)


def add_toc(p):
    r = p.add_run()
    a = OxmlElement('w:fldChar'); a.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = 'TOC \\o "1-3" \\h \\z \\u'
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = 'Para ver el indice: clic derecho sobre este texto > Actualizar campos.'
    c = OxmlElement('w:fldChar'); c.set(qn('w:fldCharType'), 'end')
    for x in (a, i, b, t, c):
        r._r.append(x)


doc = Document()
nm = doc.styles['Normal']; nm.font.name = 'Arial'; nm.font.size = Pt(12); nm.font.color.rgb = NEGRO
nm.element.rPr.rFonts.set(qn('w:cs'), 'Arial')
nm.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
nm.paragraph_format.space_after = Pt(12); nm.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def est(n, s):
    st = doc.styles[n]; st.font.name = 'Arial'; st.font.size = Pt(s); st.font.bold = True; st.font.color.rgb = NEGRO
    st.element.rPr.rFonts.set(qn('w:cs'), 'Arial')


est('Heading 1', 14); est('Heading 2', 13); est('Heading 3', 13)
doc.styles['Heading 1'].paragraph_format.space_before = Pt(18); doc.styles['Heading 1'].paragraph_format.space_after = Pt(18)
for h in ('Heading 2', 'Heading 3'):
    doc.styles[h].paragraph_format.space_before = Pt(12); doc.styles[h].paragraph_format.space_after = Pt(12)

s0 = doc.sections[0]
s0.top_margin = Cm(2.5); s0.bottom_margin = Cm(2.5); s0.left_margin = Cm(3.5); s0.right_margin = Cm(2.5)


def par(t, primero=False):
    p = doc.add_paragraph(); p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(12); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if not primero:
        p.paragraph_format.first_line_indent = Cm(1)
    _arial(p.add_run(t), 12); return p


def tit(t, n=1):
    h = doc.add_heading(level=n); _arial(h.add_run(t), 14 if n == 1 else 13, bold=True); return h


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(6); _arial(p.add_run(it), 12)


def center(t, s, bold=False, sa=12):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(sa); p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _arial(p.add_run(t), s, bold=bold); return p


def tabla(cap, heads, rows):
    c = doc.add_paragraph(); c.paragraph_format.space_before = Pt(6); c.paragraph_format.space_after = Pt(4)
    _arial(c.add_run(cap), 10, bold=True)
    t = doc.add_table(rows=1, cols=len(heads)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(heads):
        _arial(t.rows[0].cells[i].paragraphs[0].add_run(h), 10, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            _arial(cells[i].paragraphs[0].add_run(v), 10)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ===== PORTADA =====
center('INSTITUTO TECNOLÓGICO', 14, bold=True, sa=2)
center('[NOMBRE DE LA UNIVERSIDAD / INSTITUTO]', 12, sa=24)
center('PROYECTO DE INVESTIGACIÓN', 14, bold=True, sa=24)
center('[Nombre de la carrera]', 12, sa=36)
center('EduFolio: desarrollo de un portafolio virtual para la gestión '
       'y resguardo del trabajo docente', 14, bold=True, sa=12)
center('Reporte de la Segunda Fase', 13, bold=True, sa=40)
center('Autor(a):', 12, bold=True, sa=2)
center('[Nombre completo del estudiante]', 12, sa=20)
center('Asesor interno:', 12, bold=True, sa=2)
center('[Nombre del asesor(a)]', 12, sa=20)
center('Titular de la Subdirección de Estudios Profesionales:', 12, bold=True, sa=2)
center('[Nombre del titular]', 12, sa=36)
center('Junio de 2026', 12, bold=True)

# ===== PRELIMINARES =====
doc.add_section(WD_SECTION.NEW_PAGE)
sp = doc.sections[1]
sp.top_margin = Cm(2.5); sp.bottom_margin = Cm(2.5); sp.left_margin = Cm(3.5); sp.right_margin = Cm(2.5)
numbering(sp, 'lowerRoman', start=1); footer_num(sp)

tit('Resumen', 1)
par('El presente proyecto de investigación documenta el análisis, diseño y '
    'desarrollo de EduFolio, una aplicación web tipo portafolio virtual orientada '
    'al docente, que reúne en un solo espacio seguro sus documentos, notas, '
    'material didáctico y tareas. El problema que motiva el estudio es la '
    'dispersión de los recursos de trabajo del profesorado en medios físicos y '
    'servicios digitales heterogéneos, lo que dificulta su organización, '
    'resguardo y consulta.', primero=True)
par('El desarrollo sigue un modelo incremental en tres fases. La primera fase '
    'construyó la base del sistema (autenticación, arquitectura, modelo de datos '
    'y estructura de secciones) y su despliegue en línea. La segunda fase, que '
    'reporta este documento, implementa la gestión completa de contenidos: las '
    'operaciones de alta, consulta, edición y baja en las cuatro secciones, así '
    'como la carga y descarga de archivos con almacenamiento seguro. El sistema '
    'se construye con PHP 8.2 y MariaDB/MySQL, aplicando una arquitectura en capas '
    'y prácticas de seguridad como el cifrado de contraseñas, las sentencias '
    'preparadas, la protección contra CSRF y el control de acceso a los archivos.')
par('Los resultados confirman que cada docente gestiona su propio portafolio de '
    'forma íntegra y segura: solo accede a su información y a sus archivos. La '
    'tercera fase abordará el refinamiento, los roles de usuario, el despliegue '
    'definitivo y las conclusiones del proyecto.')
p = doc.add_paragraph(); _arial(p.add_run('Palabras clave: '), 12, bold=True)
_arial(p.add_run('portafolio docente, aplicación web, CRUD, carga de archivos, '
                 'PHP, MySQL, seguridad web.'), 12)

tit('Abstract', 1)
par('This research project documents the analysis, design and development of '
    'EduFolio, a web-based virtual portfolio for teachers that gathers their '
    'documents, notes, teaching materials and tasks in a single secure space. '
    'The problem addressed is the dispersion of teachers’ working resources '
    'across physical media and heterogeneous digital services, which hinders '
    'their organization, safekeeping and retrieval.', primero=True)
par('Development follows an incremental three-phase model. The first phase built '
    'the core of the system (authentication, architecture, data model and section '
    'structure) and its online deployment. The second phase, reported here, '
    'implements full content management: create, read, update and delete '
    'operations across the four sections, together with file upload and download '
    'using secure storage. The system is built with PHP 8.2 and MariaDB/MySQL, '
    'applying a layered architecture and security practices such as password '
    'hashing, prepared statements, CSRF protection and file access control.')
par('The results confirm that each teacher manages their own portfolio fully and '
    'securely, accessing only their own information and files. The third phase '
    'will address refinement, user roles, final deployment and the project '
    'conclusions.')
p = doc.add_paragraph(); _arial(p.add_run('Keywords: '), 12, bold=True)
_arial(p.add_run('teaching portfolio, web application, CRUD, file upload, PHP, '
                 'MySQL, web security.'), 12)

tit('Índice de contenido', 1)
pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
_arial(pp.add_run('El siguiente índice se genera automáticamente a partir de los títulos del documento.'), 12)
add_toc(doc.add_paragraph())

# ===== CUERPO =====
doc.add_section(WD_SECTION.NEW_PAGE)
sc = doc.sections[2]
sc.top_margin = Cm(2.5); sc.bottom_margin = Cm(2.5); sc.left_margin = Cm(3.5); sc.right_margin = Cm(2.5)
numbering(sc, 'decimal', start=1); footer_num(sc)

tit('1. Antecedentes', 1)
par('El proyecto se ubica en el área de la ingeniería de software y, de manera '
    'específica, en la tecnología educativa, en la línea de desarrollo de '
    'aplicaciones web para el apoyo de la labor docente. La incorporación de las '
    'Tecnologías de la Información y la Comunicación (TIC) ha transformado la '
    'manera en que los profesores planifican, organizan y resguardan su trabajo '
    '(Cabero Almenara, 2007; Coll, 2008). Las plataformas de gestión del '
    'aprendizaje (LMS) como Moodle y Google Classroom popularizaron el espacio '
    'digital centralizado, pero están orientadas al curso y a la relación con el '
    'alumnado, no al docente como gestor de su propio acervo. El portafolio '
    'docente, en su versión digital, favorece la organización y la reflexión '
    'profesional (Barberà y de Martín, 2009), brecha que este proyecto atiende '
    'con una herramienta ligera y centrada en el profesor.', primero=True)

tit('2. Planteamiento del problema', 1)
par('En la práctica, muchos docentes administran su trabajo mediante recursos '
    'dispersos: archivos en distintas computadoras y memorias, documentos '
    'impresos, correos y servicios no integrados. Esto genera pérdida de tiempo, '
    'riesgo de extravío y dificultad para reutilizar recursos. De ello surge la '
    'pregunta principal: ¿de qué manera una aplicación web tipo portafolio '
    'virtual puede apoyar al docente en la organización, el resguardo y la '
    'consulta centralizada de sus documentos, notas, material didáctico y '
    'tareas?', primero=True)

tit('3. Objetivos', 1)
tit('3.1 Objetivo general', 2)
par('Desarrollar una aplicación web tipo portafolio virtual para docentes que '
    'permita registrar usuarios y gestionar de forma centralizada y segura sus '
    'documentos, notas, material didáctico y tareas, implementada con tecnologías '
    'web de libre acceso.', primero=True)
tit('3.2 Objetivos específicos', 2)
bullets([
    'Analizar las necesidades de gestión de información del docente.',
    'Diseñar la arquitectura y el modelo de base de datos del sistema.',
    'Implementar la autenticación segura de usuarios.',
    'Desarrollar la gestión completa de contenidos (alta, consulta, edición y '
    'baja) y la carga de archivos en cada sección.',
    'Desplegar la aplicación y validar su funcionamiento.',
])

tit('4. Delimitación y alcances', 1)
par('El sistema está dirigido a docentes; cada usuario gestiona únicamente su '
    'propio portafolio. Las secciones consideradas son Documentos, Notas, '
    'Material didáctico y Tareas. La autenticación se realiza con correo y '
    'contraseña. El desarrollo se entrega de forma incremental en tres fases; '
    'este documento corresponde a la segunda. No forma parte del alcance la '
    'interacción directa con alumnos ni la calificación de actividades, a '
    'diferencia de un LMS completo.', primero=True)

tit('5. Hipótesis', 1)
par('H1. El desarrollo de una aplicación web tipo portafolio virtual permitirá '
    'al docente centralizar y resguardar sus recursos, mejorando la organización '
    'de su trabajo respecto al uso de medios dispersos.', primero=True)
par('H2. El uso de tecnologías web de libre acceso y de prácticas de seguridad '
    'estándar permitirá obtener una solución funcional, segura y de bajo costo.')

tit('6. Justificación', 1)
par('El proyecto se justifica por su utilidad práctica para resolver un problema '
    'cotidiano del docente y por el aprovechamiento de tecnologías accesibles. Su '
    'importancia es teórica (articula tecnología educativa e ingeniería de '
    'software), metodológica (modelo incremental replicable), práctica (reduce '
    'tiempo de búsqueda y riesgo de pérdida) y social (solución de bajo costo y '
    'replicable).', primero=True)

tit('7. Marco teórico', 1)
tit('7.1 El portafolio docente y el portafolio digital', 2)
par('El portafolio docente es una colección organizada de evidencias del trabajo '
    'del profesor; su versión digital facilita el almacenamiento, la '
    'actualización y la consulta (Barberà y de Martín, 2009).', primero=True)
tit('7.2 Sistemas de gestión del aprendizaje (LMS)', 2)
par('Los LMS integran cursos, contenidos y usuarios. Moodle y Google Classroom '
    'son referentes, pero su orientación al curso los hace complejos para el uso '
    'individual del docente (Area Moreira, 2009).', primero=True)
tit('7.3 Desarrollo de aplicaciones web y bases de datos', 2)
par('La aplicación se construye con PHP del lado del servidor (The PHP Group, '
    '2024) y MariaDB/MySQL como gestor relacional (Date, 2001; Silberschatz '
    'et al., 2011). El acceso a datos usa PDO con sentencias preparadas, lo que '
    'separa el código de los datos.', primero=True)
tit('7.4 Arquitectura de software en capas', 2)
par('La arquitectura en capas separa presentación, lógica de dominio y acceso a '
    'datos (Fowler, 2002), lo que favorece el mantenimiento y la extensión '
    '(Sommerville, 2011; Pressman, 2010). EduFolio adopta este modelo con un '
    'controlador frontal.', primero=True)
tit('7.5 Seguridad en aplicaciones web y manejo de archivos', 2)
par('Siguiendo las recomendaciones de la OWASP Foundation (2021), se adoptan: '
    'cifrado de contraseñas con funciones de hash, sentencias preparadas contra '
    'inyección SQL, escape de la salida contra XSS, tokens CSRF en formularios y '
    'manejo seguro de la sesión. En la carga de archivos se valida tipo y tamaño, '
    'se almacenan fuera del directorio público y se sirven mediante un script que '
    'verifica la propiedad del archivo, evitando el acceso directo por URL.',
    primero=True)

tit('8. Metodología', 1)
tit('8.1 Tipo de investigación', 2)
par('Investigación aplicada de carácter tecnológico, cuyo producto principal es '
    'el desarrollo de software, con alcance descriptivo en el análisis y '
    'propositivo en la solución (Hernández Sampieri et al., 2014).', primero=True)
tit('8.2 Metodología de desarrollo de software', 2)
par('Modelo incremental en tres fases entregables, cada una con análisis, '
    'diseño, implementación y pruebas (Pressman, 2010):', primero=True)
bullets([
    'Fase 1 — Base del sistema: arquitectura, datos, autenticación, estructura y '
    'despliegue inicial.',
    'Fase 2 — Gestión de contenidos: operaciones CRUD y carga de archivos en cada '
    'sección.',
    'Fase 3 — Refinamiento y cierre: roles de usuario, ajustes, despliegue '
    'definitivo y resultados.',
])
tit('8.3 Arquitectura del sistema', 2)
par('El sistema se organiza en capas. La carpeta de lógica permanece fuera del '
    'alcance web y solo la carpeta pública se expone al navegador; un controlador '
    'frontal dirige las peticiones. La Tabla 1 resume las capas.', primero=True)
tabla('Tabla 1. Capas de la arquitectura de EduFolio',
      ['Capa', 'Responsabilidad'],
      [['Enrutamiento', 'Dirige toda petición hacia el directorio público.'],
       ['Presentación', 'Vistas y páginas que reciben la petición.'],
       ['Lógica / Dominio', 'Reglas de negocio: autenticación, validaciones, CRUD.'],
       ['Acceso a datos', 'Conexión y consultas con PDO.'],
       ['Configuración', 'Parámetros del entorno y sesión.'],
       ['Datos', 'Base de datos relacional y almacenamiento de archivos.']])
tit('8.4 Tecnologías utilizadas', 2)
tabla('Tabla 2. Tecnologías empleadas en el desarrollo',
      ['Componente', 'Tecnología'],
      [['Lenguaje de servidor', 'PHP 8.2 (compatible con PHP 7.2+)'],
       ['Base de datos', 'MariaDB / MySQL'],
       ['Acceso a datos', 'PDO con sentencias preparadas'],
       ['Interfaz', 'HTML5, CSS3 y JavaScript'],
       ['Iconografía y tipografía', 'Bootstrap Icons y Plus Jakarta Sans'],
       ['Entorno local', 'XAMPP (Apache, MariaDB, PHP)'],
       ['Despliegue', 'InfinityFree (PHP 8 y MySQL)']])
tit('8.5 Diseño de la base de datos', 2)
par('El modelo es relacional: una tabla de usuarios y cuatro tablas de secciones '
    'relacionadas mediante claves foráneas con borrado en cascada (Tabla 3).',
    primero=True)
tabla('Tabla 3. Entidades de la base de datos',
      ['Tabla', 'Descripción'],
      [['usuarios', 'Cuenta del docente y contraseña cifrada.'],
       ['documentos', 'Archivos institucionales del usuario.'],
       ['notas', 'Apuntes de texto del docente.'],
       ['materiales', 'Material didáctico organizado por materia.'],
       ['tareas', 'Actividades con fecha de entrega y estado.']])

tit('9. Desarrollo del sistema', 1)
par('Esta sección describe lo construido y verificado hasta la segunda fase.',
    primero=True)
tit('9.1 Autenticación y estructura (Fase 1)', 2)
par('Se implementó el registro de docentes, el inicio y el cierre de sesión, con '
    'contraseñas cifradas mediante bcrypt, y se construyó el panel principal con '
    'la estructura navegable de las cuatro secciones. La aplicación quedó '
    'desplegada y accesible en línea.', primero=True)
tit('9.2 Gestión de contenidos: operaciones CRUD (Fase 2)', 2)
par('Cada sección dispone de las operaciones de alta, consulta, edición y baja, '
    'siempre acotadas al usuario en sesión: un docente solo puede ver y modificar '
    'sus propios registros. La Tabla 4 resume las operaciones por sección.',
    primero=True)
tabla('Tabla 4. Operaciones de gestión por sección',
      ['Sección', 'Crear', 'Consultar', 'Editar', 'Eliminar'],
      [['Notas', 'Sí', 'Sí', 'Sí', 'Sí'],
       ['Documentos', 'Sí', 'Sí', 'Sí', 'Sí'],
       ['Material didáctico', 'Sí', 'Sí', 'Sí', 'Sí'],
       ['Tareas', 'Sí', 'Sí', 'Sí', 'Sí']])
par('Las notas gestionan título y contenido; las tareas incluyen además fecha de '
    'entrega y estado (pendiente, en progreso o completada), con resaltado de las '
    'tareas vencidas. El panel muestra un contador de elementos por sección.')
tit('9.3 Carga y descarga de archivos (Fase 2)', 2)
par('Las secciones de Documentos y Material permiten subir archivos. Al cargar, '
    'el sistema valida el tipo (PDF, Word, Excel, PowerPoint, imágenes, texto y '
    'comprimidos) y el tamaño (máximo 10 MB), genera un nombre único y guarda el '
    'archivo en una carpeta propia del usuario, ubicada fuera del directorio '
    'público. La descarga se realiza mediante un controlador que verifica que el '
    'archivo pertenezca al usuario antes de entregarlo, de modo que no es posible '
    'acceder a los archivos de otros por URL.', primero=True)
tit('9.4 Seguridad implementada', 2)
par('Se incorporaron las prácticas descritas en el marco teórico: cifrado de '
    'contraseñas, sentencias preparadas, escape de la salida, tokens CSRF en los '
    'formularios, manejo seguro de la sesión, almacenamiento de archivos fuera '
    'del alcance web y descarga controlada por propietario (OWASP Foundation, '
    '2021).', primero=True)

tit('10. Conclusiones', 1)
par('La segunda fase cumplió su objetivo: las cuatro secciones del portafolio '
    'cuentan con gestión completa de contenidos y, donde corresponde, con carga y '
    'descarga segura de archivos. Con ello, el sistema responde de forma efectiva '
    'a la pregunta principal de investigación y respalda la hipótesis sobre la '
    'mejora en la organización del trabajo docente (H1), manteniendo la '
    'viabilidad técnica y de bajo costo de la solución (H2).', primero=True)
par('El sistema resguarda la información de cada docente de manera aislada y '
    'segura. Queda pendiente, para la tercera fase, el refinamiento, la '
    'incorporación de roles de usuario, el despliegue definitivo y la '
    'documentación de los resultados finales.')

tit('11. Recomendaciones, líneas abiertas o trabajos futuros', 1)
bullets([
    'Fase 3: incorporar roles (administrador y docente), refinar la interfaz y '
    'realizar el despliegue definitivo con pruebas finales.',
    'Añadir validación del tipo real de archivo (MIME), edición de perfil y '
    'cambio de contraseña.',
    'Incorporar búsqueda y filtrado de contenidos, categorización por materia o '
    'ciclo escolar y exportación del portafolio.',
])

tit('Referencias', 1)
refs = [
    'Area Moreira, M. (2009). Introducción a la tecnología educativa. Universidad de La Laguna.',
    'Barberà, E., & de Martín, E. (2009). Portfolio electrónico: aprender a evaluar el aprendizaje. Editorial UOC.',
    'Cabero Almenara, J. (2007). Nuevas tecnologías aplicadas a la educación. McGraw-Hill.',
    'Coll, C. (2008). Aprender y enseñar con las TIC: expectativas, realidad y potencialidades. Boletín de la Institución Libre de Enseñanza, 72, 17-40.',
    'Date, C. J. (2001). Introducción a los sistemas de bases de datos (7.ª ed.). Pearson Educación.',
    'Fowler, M. (2002). Patterns of enterprise application architecture. Addison-Wesley.',
    'Hernández Sampieri, R., Fernández Collado, C., & Baptista Lucio, P. (2014). Metodología de la investigación (6.ª ed.). McGraw-Hill.',
    'Nielsen, J. (1994). Usability engineering. Morgan Kaufmann.',
    'OWASP Foundation. (2021). OWASP Top 10. https://owasp.org/Top10/',
    'Pressman, R. S. (2010). Ingeniería del software: un enfoque práctico (7.ª ed.). McGraw-Hill.',
    'Silberschatz, A., Korth, H. F., & Sudarshan, S. (2011). Database system concepts (6.ª ed.). McGraw-Hill.',
    'Sommerville, I. (2011). Ingeniería de software (9.ª ed.). Pearson Educación.',
    'The PHP Group. (2024). PHP manual. https://www.php.net/manual/es/',
    'UNESCO. (2019). Marco de competencias de los docentes en materia de TIC (versión 3). UNESCO.',
]
for ref in sorted(refs):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE; p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.left_indent = Cm(1); p.paragraph_format.first_line_indent = Cm(-1)
    _arial(p.add_run(ref), 12)

salida = __file__.rsplit('\\', 1)[0] + '\\EduFolio_Documentacion_Fase2.docx'
doc.save(salida)
print('OK ->', salida)
